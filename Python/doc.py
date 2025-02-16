from docx import Document

# Create a new Word Document
doc = Document()

# Title
doc.add_heading('Practical 1 - Software Engineering', level=1)

# Introduction Section
doc.add_heading('INTRODUCTION:', level=2)

doc.add_heading('Aim:', level=3)
doc.add_paragraph(
    "To design and develop software that helps users create, format, and manage notes efficiently while reading or "
    "researching, incorporating features like AI assistance, text formatting, and exporting to a Word document."
)

doc.add_heading('Innovative Features of the System:', level=3)
doc.add_paragraph(
    "1. AI-Powered Assistance: Provides instant YouTube suggestions and answers questions using AI-based chat.\n"
    "2. Text Formatting Options: Supports bold, headings (H1, H2), bullet points, and text alignment.\n"
    "3. Pop-up Interaction: Activates only when users select text, avoiding unnecessary interruptions.\n"
    "4. Instant Search: Allows users to search saved notes efficiently.\n"
    "5. Screenshot Integration: Capture and embed screenshots directly into notes."
)

doc.add_heading('Scope of the System:', level=3)
doc.add_paragraph(
    "1. Ideal for students, researchers, and professionals for efficient note-taking.\n"
    "2. Useful in online study environments, classrooms, or during personal research.\n"
    "3. Saves time by providing AI-based recommendations and formatting tools.\n"
    "4. Encourages organized note management, suitable for future references or exporting to editable Word files."
)

# Description Section
doc.add_heading('DESCRIPTION:', level=2)

doc.add_heading('Main Modules of the System:', level=3)
doc.add_paragraph(
    "1. Text Selection and Note Compilation: Appends selected text to notes and formats it using predefined tools.\n"
    "2. AI Assistance: Suggests relevant videos or provides answers to user queries via chat.\n"
    "3. Search and Organize Notes: Enables users to search through saved notes instantly.\n"
    "4. Export Module: Compiles notes into a well-structured Word document.\n"
    "5. Screenshot Integration: Captures and attaches important screenshots to the notes."
)

doc.add_heading('Requirement of the System:', level=3)
doc.add_paragraph(
    "1. Client Requirements:\n"
    "   - Simple and user-friendly interface.\n"
    "   - Responsive system with minimal lag for operations.\n"
    "   - Compatibility with Windows operating system.\n\n"
    "2. System Requirements:\n"
    "   - Functional Requirements:\n"
    "       - Text selection and formatting.\n"
    "       - AI integration for queries and suggestions.\n"
    "       - Save notes in Word format.\n"
    "   - Non-Functional Requirements:\n"
    "       - System should have a fast response time (within 2 seconds for actions).\n"
    "       - Ensure compatibility with Windows 10 and above."
)

doc.add_heading('Feasibility Studies:', level=3)
doc.add_paragraph(
    "1. Technical Feasibility:\n"
    "   - The system is feasible to develop using Python and Tkinter for UI, with additional libraries like pyautogui "
    "for screenshots and openai for AI chat.\n"
    "   - The hardware requirements (moderate RAM and processor) are met by most modern systems.\n\n"
    "2. Operational Feasibility:\n"
    "   - Users with basic computer knowledge can operate the system.\n"
    "   - The design ensures minimal training time for users to adapt to the tool."
)

doc.add_heading('Software Requirements:', level=3)
doc.add_paragraph(
    "1. Python 3.12 or later.\n"
    "2. Libraries: Tkinter, PyAutoGUI, OpenAI API, Pillow, Docx.\n"
    "3. Microsoft Word or any Word file viewer for exported files."
)

doc.add_heading('Technology Used for Development:', level=3)
doc.add_paragraph(
    "1. Frontend Development: Python (Tkinter for GUI).\n"
    "2. Backend Development: Python for logic, AI integration using OpenAI API.\n"
    "3. Database (optional): SQLite for saving notes locally."
)

doc.add_heading('Operating System:', level=3)
doc.add_paragraph("Windows 10 or higher (64-bit).")

# Save the document
file_path = "/mnt/data/Practical_1_Software_Engineering.docx"
doc.save(file_path)

file_path
